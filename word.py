import os
import os.path
import docx
import shutil
import subprocess
from tkinter import messagebox
from helper_methods import greek_accent_remover

class WordCreator():
    def __init__(self, main):
        self.main = main

    def create_doc(self):
        self.name_cap = greek_accent_remover(self.main.name).upper()
        self.doc_filename = self.main.surname + ' ' + self.name_cap + '.doc'
        self.doc = docx.Document()
        self.title = f'{self.main.reason} του {self.main.id_number} δελτίου ταυτότητας{self.other_docs_title} με στοιχεία: {self.main.surname} {self.main.name}.'

        self.doc.add_paragraph(self.title)
        self.doc.add_paragraph('')
        self.doc.add_paragraph('ΣΧΕΤ.: α) Η 8200/0 – 469448 από 05/10/2014 εγκύκλιος διαταγή.')
        if self.main.card == 1:
            self.doc.add_paragraph(f'            β) Η 71675/12/439946 από 07/04/2012 διαταγή.')
        self.doc.add_paragraph(f'            {self.card_text}) Το {self.main.protocol_num}{self.from_prot_date}{self.main.protocol_date_string} {self.doc_type} του {self.office_type_text} {self.main.office_article} {self.main.office_name}.')
        self.doc.add_paragraph('')
        self.doc.add_paragraph(f'  Σας διαβιβάζουμε το ανωτέρω ({self.card_text}) σχετικό και παρακαλούμε όπως προβείτε στην άμεση ακύρωση του εν θέματι δελτίου ταυτότητας, λόγω {self.reason_text} του, καταχωρώντας τις προβλεπόμενες μεταβολές στην κεντρική εφαρμογή ταυτοτήτων σύμφωνα με την ανωτέρω (α) εγκύκλιο.')
        if self.main.card == 1:
            self.doc.add_paragraph(f'  Εφιστούμε την προσοχή σας για τη σάρωση της καρτέλας (αίτηση - φωτογραφία) βάσει της (β) σχετικής, προ της καταστροφής των δικαιολογητικών.')
        self.doc.add_paragraph(self.other_docs_par)
        self.doc.save(self.doc_filename)
        
        self.move_files()

        subprocess.Popen(self.doc_destination, shell=True)

    def create_text_variables(self):
        if self.main.reason == 'Απώλεια':
            self.reason_text = 'απώλειάς'
        elif self.main.reason == 'Κλοπή':
            self.reason_text = 'κλοπής'
        else:
            self.reason_text = 'κατάσχεσής'

        self.doc_type = 'έγγραφο'

        if self.main.office_type == 1:
            self.office_type_text = 'Προξενικού Γραφείου της Πρεσβείας της Ελλάδος'
        elif self.main.office_type == 2:
            self.office_type_text = 'Γενικού Προξενείου της Ελλάδος'
        elif self.main.office_type == 3:
            self.office_type_text = 'Επιτίμου Γενικού Προξενείου της Ελλάδος'
        elif self.main.office_type == 4:
            self.office_type_text = 'Άμισθου Γενικού Προξενείου της Ελλάδος'
        else:
            self.office_type_text = 'Κεντρικού Λιμεναρχείου'
            self.doc_type = 'σήμα'

        if self.main.card == 0:
            self.card_text = 'β'
        else:
            self.card_text = 'γ'

        if self.main.other_doc_passport == 0 and self.main.other_doc_driver == 0:
            self.other_docs_title = ''
            self.other_docs_par =''
        elif self.main.other_doc_passport == 1 and self.main.other_doc_driver == 0:
            self.other_docs_title = ' και διαβατηρίου'
            self.other_docs_par = '  Στη Διεύθυνση Διαβατηρίων και Εγγράφων Ασφαλείας/Α.Ε.Α., το παρόν κοινοποιείται για ενημέρωση και τις δικές της περαιτέρω ενέργειες, ως προς την ακύρωση του διαβατηρίου.'
        elif self.main.other_doc_passport == 0 and self.main.other_doc_driver == 1:
            self.other_docs_title = ' και αδείας ικανότητας οδήγησης'
            self.other_docs_par = '  Στη Διεύθυνση Διαβατηρίων και Εγγράφων Ασφαλείας/Α.Ε.Α., το παρόν κοινοποιείται για ενημέρωση και τις τυχόν δικές της ενέργειες.'
        else:
            self.other_docs_title = ', διαβατηρίου και αδείας ικανότητας οδήγησης'
            self.other_docs_par = '  Στη Διεύθυνση Διαβατηρίων και Εγγράφων Ασφαλείας/Α.Ε.Α., το παρόν κοινοποιείται για ενημέρωση και τις τυχόν δικές της ενέργειες.'

        if self.main.protocol_date:
            self.from_prot_date = ' από '
        else:
            self.from_prot_date = ''

        self.create_doc()

    
    def move_files(self):
        self.desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
        self.pdf_filename = self.doc_filename.split(".")[0] + '.pdf'
        self.office_folder = os.path.join(self.main.folder_selected, self.main.office_name)
        if self.main.office_type == 5:
            self.office_folder = os.path.join(self.main.folder_selected, "Λιμεναρχεία")
        os.makedirs(self.office_folder, exist_ok=True)
        if not os.path.isfile(os.path.join(self.office_folder, self.doc_filename)):
            shutil.move(self.doc_filename, self.office_folder)
            self.doc_destination = os.path.join(self.office_folder, self.doc_filename)
            while True:
                try:
                    os.rename(os.path.join(self.desktop, self.pdf_filename), os.path.join(self.office_folder, self.pdf_filename))
                    break
                except PermissionError:
                    messagebox.showerror('Σφάλμα!', f'Παρακαλώ κλείστε το αρχείο {self.pdf_filename} προκειμένου ολοκληρωθεί η αρχειοθέτηση!')
                except FileExistsError:
                    messagebox.showerror('Σφάλμα!', f'Το αρχείο με τίτλο {self.pdf_filename} υπάρχει ήδη στον φάκελο {self.main.office_name}!\nΠαρακαλώ μετακινήσετε το αρχείο χειροκίνητα!')
                    break
                except OSError:
                    shutil.move(os.path.join(self.desktop, self.pdf_filename), os.path.join(self.office_folder, self.pdf_filename))
                    break
        else:
            os.makedirs(os.path.join(self.desktop, 'Απώλειες-Κλοπές'), exist_ok=True)
            shutil.move(self.doc_filename, os.path.join(self.desktop, os.path.join('Απώλειες-Κλοπές', self.doc_filename)))
            self.doc_destination = os.path.join(self.desktop, os.path.join('Απώλειες-Κλοπές', self.doc_filename))
            messagebox.showerror('Σφάλμα!', f'Το αρχείο με τίτλο {self.doc_filename} υπάρχει ήδη στον φάκελο {self.main.office_name}!\nΤο αρχείο μεταφέρθηκε στην επιφάνεια εργασίας προκειμένου να το μετακινήσετε χειροκίνητα!')
